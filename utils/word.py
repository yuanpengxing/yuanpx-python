# -*- coding: UTF-8 -*-
# author: yuanpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


class DocxUtil:
    @classmethod
    def add_paragraph(cls, document, text, bold=False, fontsize=12, alignment=0):
        # document = Document()
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if alignment == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if alignment == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(fontsize)
        run.font.name = '宋体'
        run.bold = bold

    @classmethod
    def add_heading(cls, document, title, level=2, fontsize=None):
        head = document.add_heading(level=level)
        head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        head.paragraph_format.space_after = Pt(15)
        run = head.add_run(title)
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if not fontsize:
            if level == 1:
                run.font.size = Pt(15)
            elif level == 2:
                run.font.size = Pt(14)
            elif level == 3:
                run.font.size = Pt(12)
        else:
            run.font.size = Pt(fontsize)

    @classmethod
    def add_picture(cls, document, imgpath):
        document.add_picture(imgpath, width=Inches(5.2))
